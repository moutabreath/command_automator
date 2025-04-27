import re
from docx import Document
import docx
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement
from docx.shared import Inches

def save_text_as_word(file_path, applicant_name, resume_text, resume_sections = []):
    applicant_name = applicant_name.replace("_", " ")
    # Create a Word document
    doc = Document()
    set_document_margins(doc)
    style = doc.styles['Normal']
    style.font.size = Pt(9)
    style.paragraph_format.line_spacing_rule = 1

    # Split the resume text into lines
    lines = resume_text.split('\n')

    # Add the resume content
    for line in lines:
        if add_applicant_name(doc, applicant_name, line):
            continue
        is_header = add_header(doc, resume_sections, line)
        if is_header:
            is_header = False
            continue
        url = extract_link(line)
        if url is not None:                    
            line = add_line_with_link(doc, line, url)            
            continue           

        if "*" in line:
            p = doc.add_paragraph()            
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            line.replace("**", "")
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(9)
            continue
        if line.startswith('-'):
            p = doc.add_paragraph(line[1:], style='ListBullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.size = Pt(9)
            continue
        p = doc.add_paragraph(line)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.size = Pt(9)

    doc.save(file_path)

def set_document_margins(doc, top=0.6, bottom=0.4, left=0.75, right=0.75):
    """
    Set the margins for the document in inches.
    :param doc: The Document object.
    :param top: Top margin in inches.
    :param bottom: Bottom margin in inches.
    :param left: Left margin in inches.
    :param right: Right margin in inches.
    """
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def add_applicant_name(doc, applicant_name:str, line: str):    
    if applicant_name == line:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 255)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        return True
    return False


def add_line_with_link(doc, line, url, p = None):
    link_name = "link"
    if 'linkedin' in line:
        link_name = "LinkedIn"
    if 'github' in line:
        link_name = get_github_project_name(url)
    line = line.replace(url, "")
    if p is None:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    p.add_run(line)
    add_link(link_name, url, p)


def get_github_project_name(url):
    # Regex pattern to find the last part of the URL after the last '/'
    pattern = r'[^/]+$'
    match = re.search(pattern, url)
    if match:
        return match.group()
    else:
        return None

def extract_link(text):
    # Regex pattern to find URLs
    pattern =  r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
    match = re.search(pattern, text)
    if match:
        return match.group()
    else:
        return None

def add_link(link_name, url, paragraph):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = link_name
    new_run.font.size = Pt(9)
    new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)

#This is only needed if you're using the builtin style above
def get_or_create_hyperlink_style(d):
    """If this document had no hyperlinks so far, the builtin
       Hyperlink style will likely be missing and we need to add it.
       There's no predefined value, different Word versions
       define it differently.
       This version is how Word 2019 defines it in the
       default theme, excluding a theme reference.
    """
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style("Default Character Font",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style("Hyperlink",
                                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style 

def add_header(doc, keywords, line: str):
    text = line.lower() 
    is_header = False       
    for keyword in keywords:
        if keyword in text:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            
            index = text.find(keyword)
            line_keyword = line[index:len(keyword) + 1] # +1 for space character
            line = line.replace(line_keyword, "", 1).strip()
            run = p.add_run(line_keyword)
            run.bold = True 
            run.font.color.rgb = RGBColor(0, 0, 255)
            
            is_header = True
            url = extract_link(line)
            if  url is not None:                    
                add_line_with_link(doc, line, url, p)
            break
    return is_header